"""Create canonical v2 copies of the Persian thesis and project report."""

from __future__ import annotations

import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def _copy_run_properties(source: Paragraph, target: Paragraph) -> None:
    if source.runs and target.runs and source.runs[0]._r.rPr is not None:
        target.runs[0]._r.insert(0, deepcopy(source.runs[0]._r.rPr))


def _replace(paragraph: Paragraph, text: str) -> None:
    run_properties = (
        deepcopy(paragraph.runs[0]._r.rPr)
        if paragraph.runs and paragraph.runs[0]._r.rPr is not None
        else None
    )
    for child in list(paragraph._p):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            paragraph._p.remove(child)
    run = paragraph.add_run(text)
    if run_properties is not None:
        run._r.insert(0, run_properties)


def _insert_after(anchor: Paragraph, text: str, template: Paragraph) -> Paragraph:
    element = OxmlElement("w:p")
    anchor._p.addnext(element)
    paragraph = Paragraph(element, anchor._parent)
    if template._p.pPr is not None:
        paragraph._p.insert(0, deepcopy(template._p.pPr))
    paragraph.add_run(text)
    _copy_run_properties(template, paragraph)
    return paragraph


def _replace_cell(cell: object, text: str) -> None:
    paragraphs = cell.paragraphs  # type: ignore[attr-defined]
    _replace(paragraphs[0], text)
    for paragraph in paragraphs[1:]:
        _replace(paragraph, "")


def _canonical_notice(
    document: Document, anchor_index: int, heading_index: int, body_index: int
) -> None:
    anchor = document.paragraphs[anchor_index]
    heading_template = document.paragraphs[heading_index]
    body_template = document.paragraphs[body_index]
    page_break = _insert_after(anchor, "", body_template)
    page_break.add_run().add_break(WD_BREAK.PAGE)
    heading = _insert_after(
        page_break,
        "یادداشت اصلاح علمی و مرجع canonical (۲۲ اوت ۲۰۲۶)",
        heading_template,
    )
    body_1 = _insert_after(
        heading,
        (
            "مرجع جاری نتایج، docs/qa_remediation_v2_report.md و اجرای فریز‌شدهٔ "
            "executable_direction_hurst_ablation_v2-20260822T165905Z است. هر ۱۲ آزمایش "
            "توسعه رد شدند؛ بهترین دقت متوازن ۰٫۵۲۰۳۵ برای no-Hurst در افق یک‌روزه بود "
            "و هیچ candidate فریز نشد."
        ),
        body_template,
    )
    body_2 = _insert_after(
        body_1,
        (
            "بازهٔ ۲۰۱۱ تا ۳۰ ژوئن ۲۰۲۳ شواهد توسعهٔ قبلاً دیده‌شده است. بازهٔ ۳ ژوئیهٔ "
            "۲۰۲۳ تا ۳۱ ژوئیهٔ ۲۰۲۶ historical_audit_previously_revealed است و در اجرای v2 "
            "بارگذاری نشد. confirmation تاریخی تازه و شواهد future out-of-sample وجود ندارد."
        ),
        body_template,
    )
    body_3 = _insert_after(
        body_2,
        (
            "اقتصاد non-overlapping و receipt/registry در v2 اصلاح شد؛ بااین‌حال مرز "
            "calibration/evaluation داخلی بر endpoint برچسب purge نشده و هر ۶۰ انتخاب outer "
            "را متاثر می‌کند. PBO/DSR نیز صرفاً diagnostic است. هیچ ادعای leakage-free، "
            "سودآوری، برتری هِرست یا parity بومی MQL5 مجاز نیست."
        ),
        body_template,
    )
    body_3.add_run().add_break(WD_BREAK.PAGE)


def update_thesis(source: Path, output: Path) -> None:
    shutil.copy2(source, output)
    document = Document(output)
    abstract = document.paragraphs[18]
    replacements = {
        221: (
            "شواهد این فصل سه طبقهٔ جدا دارد: development قبلاً دیده‌شده تا ۳۰ ژوئن "
            "۲۰۲۳؛ historical_audit_previously_revealed از ۳ ژوئیهٔ ۲۰۲۳ تا ۳۱ ژوئیهٔ "
            "۲۰۲۶؛ و future out-of-sample که هنوز وجود ندارد. اجرای canonical v2 فقط "
            "development را بارگذاری کرد و هیچ confirmation یا audit تازه‌ای انجام نشد."
        ),
        222: (
            "خانوادهٔ مستقل v2 سه بازوی no-Hurst، DFA Hurst و robust Hurst را در چهار افق "
            "با بودجهٔ ثابت ۱۲ trial آزمود. این خانواده پس از اجرا immutable است. ممیزی "
            "پس از freeze نشت endpoint در مرز calibration/evaluation داخلی را یافت؛ بنابراین "
            "metrics آن برای promotion یا ادعای leakage-free معتبر نیست و rerun همان شناسه "
            "ممنوع است."
        ),
        234: "۵-۵. نتیجهٔ ثبت‌شدهٔ v2 و محدودیت اعتبار آن",
        236: (
            "runner هر ۱۲ trial را رد کرد. بهترین BA نقطه‌ای no-Hurst/H1 برابر ۰٫۵۲۰۳۵ و "
            "F1 کلان آن ۰٫۵۱۶۱۸ بود؛ بازده خالص ۵ و ۱۰ bps هر دو منفی شد. robust-Hurst/H5 "
            "بازده نقطه‌ای مثبت داشت، اما gateهای طبقه‌بندی، عدم‌قطعیت، calibration، PBO/DSR "
            "و QA را پاس نکرد. با توجه به leakage انتخاب calibration و نبود paired CI مستقیم، "
            "این اعداد اثبات رد یا تأیید ارزش افزودهٔ هِرست نیستند؛ تنها نشان می‌دهند هیچ "
            "candidate مجاز به promotion نشد."
        ),
        239: (
            "ممیزی مستقل receipt دارای ۱۱۰ خروجی، ۲۱ منبع runtime و registry دارای ۱۲ رکورد "
            "را معتبر یافت؛ تست‌های نرم‌افزاری نیز پاس شدند. بااین‌حال هر ۶۰ انتخاب outer در "
            "مرز calibration/evaluation بر executable_label_end_index purge نشده‌اند و sigmoid "
            "در ۲۶ مورد انتخاب شده است. recommendation نهایی برای مدل و protocol: REJECT."
        ),
        246: (
            "یافتهٔ قابل دفاع این است که هیچ candidate طبق runner فریز نشد و هیچ ادعای "
            "سودآوری یا برتری هِرست مجاز نیست. به‌علت leakage انتخاب calibration، v2 شاهد "
            "استنباطی leakage-free برای عملکرد بازار نیست. داده‌های تاریخی V3 نیز قبلاً "
            "آشکار شده‌اند و فقط زمینهٔ تاریخی‌اند، نه confirmation یا آیندهٔ واقعی."
        ),
        248: (
            "نقطهٔ قوت فعلی، ثبت شفاف شکست و حفظ immutability است: جداسازی فیزیکی داده، "
            "schedule اقتصادی non-overlapping، receipt کامل، registry hash-chained و quarantine "
            "run ناقص پیاده شده‌اند. کنترل مرز calibration، common-calendar، multiplicity و "
            "event-boundary-safe PBO/DSR باید در خانواده‌ای تازه پیش‌ثبت شوند."
        ),
        251: (
            "۲. دورهٔ ۲۰۲۳ تا ۲۰۲۶ قبلاً مشاهده شده و historical_audit_previously_revealed "
            "است؛ در v2 دوباره بارگذاری نشد و آزمون بکر محسوب نمی‌شود."
        ),
        252: (
            "۳. runner هر ۱۲ trial را رد کرد، اما leakage مرز calibration/evaluation مانع "
            "تفسیر metrics به‌عنوان آزمون معتبر ارزش افزودهٔ هِرست است."
        ),
        253: (
            "۴. CIهای طبقه‌بندی familywise-adjusted نیستند، foldها بین افق‌ها common-calendar "
            "کامل ندارند و PBO/DSR در مرز eventهای چندروزه محافظت نشده‌اند."
        ),
        263: (
            "این پایان‌نامه چارچوبی قابل ممیزی برای آزمون چندافقی XAUUSD ارائه می‌کند. "
            "اصلاح v2 دادهٔ development را فیزیکی جدا، اقتصاد را non-overlapping و تمام "
            "خروجی‌ها را receipt-hashed کرد. ممیزی مستقل در عین حال نشت مرز calibration داخلی "
            "را کشف کرد؛ ازاین‌رو شفافیت نتیجه و جلوگیری از دست‌کاری پس از freeze، نه ادعای "
            "leakage-free بودن run، دستاورد قابل دفاع است."
        ),
        265: (
            "در v2 هر ۱۲ trial رد و هیچ candidate فریز نشد. بهترین BA نقطه‌ای ۰٫۵۲۰۳۵ برای "
            "no-Hurst/H1 بود، ولی calibration-boundary leakage و سایر محدودیت‌های آماری اجازهٔ "
            "استنباط معتبر دربارهٔ مهارت، ارزش افزودهٔ هِرست یا سودآوری را نمی‌دهند. confirmation، "
            "audit جدید، MT5 candidate run و future OOS وجود ندارد."
        ),
        279: (
            "۱. طراحی خانواده‌ای تازه و پیش‌ثبت‌شده که مرز calibration/evaluation را بر endpoint "
            "برچسب purge کند، foldهای common-calendar داشته باشد، familywise inference و paired "
            "Hurst-vs-no-Hurst را اعمال و PBO/DSR را با مرز event سازگار کند؛ v2 نباید rerun یا "
            "دست‌کاری شود."
        ),
    }
    target_paragraphs = {index: document.paragraphs[index] for index in replacements}
    _canonical_notice(document, anchor_index=8, heading_index=17, body_index=18)
    _replace(
        abstract,
        (
            "چکیده — این پایان‌نامه نقش ویژگی‌های رژیمی مبتنی بر نمای هِرست را در "
            "پیش‌بینی چندافقی جهت روزانهٔ XAUUSD بررسی می‌کند. خانوادهٔ مستقل و "
            "پیش‌ثبت‌شدهٔ v2 شامل سه بازوی بدون هِرست، هِرست DFA و رژیم هِرست مقاوم در "
            "افق‌های ۱، ۵، ۱۰ و ۲۰ روز بود. هر ۱۲ آزمایش روی دادهٔ توسعهٔ قبلاً "
            "دیده‌شدهٔ ۲۰۱۱ تا ژوئن ۲۰۲۳ رد شدند. بهترین نتیجه متعلق به no-Hurst/H1 با "
            "دقت متوازن ۰٫۵۲۰۳۵، F1 کلان ۰٫۵۱۶۱۸ و بازهٔ اطمینان اسمی ۹۵٪ برابر "
            "[۰٫۵۰۱۸۹، ۰٫۵۳۸۳۷] بود؛ بازده خالص آن در هزینه‌های ۵ و ۱۰ bps منفی شد. "
            "هیچ candidate فریز نشد، audit تاریخی ۲۰۲۳ تا ۲۰۲۶ دوباره اجرا نشد و شواهد "
            "آیندهٔ واقعی وجود ندارد. بنابراین داده‌های فعلی شواهد پایدار برای ارزش "
            "افزودهٔ هِرست یا سودآوری فراهم نمی‌کنند. leakage مرز calibration/evaluation، "
            "common-calendar، کنترل چندآزمونی و PBO/DSR وابسته به مرز event صریحاً برای "
            "خانوادهٔ بعدی deferred شده‌اند."
        ),
    )
    for index, text in replacements.items():
        _replace(target_paragraphs[index], text)

    ablation = document.tables[18]
    ablation_values = [
        ["بدون هرست", "۱، ۵، ۱۰، ۲۰", "H1", "۰٫۵۲۰۳۵", "۰٫۵۱۶۱۸", "رد؛ evidence دارای leakage"],
        ["هرست DFA فعلی", "۱، ۵، ۱۰، ۲۰", "H1", "۰٫۵۰۵۳۷", "۰٫۴۷۱۹۳", "رد؛ بدون paired CI"],
        ["رژیم هرست مقاوم", "۱، ۵، ۱۰، ۲۰", "H20", "۰٫۵۱۲۳۹", "۰٫۴۶۵۴۶", "رد؛ بدون paired CI"],
    ]
    for row, values in zip(ablation.rows[1:], ablation_values, strict=True):
        for cell, value in zip(row.cells, values, strict=True):
            _replace_cell(cell, value)

    qa_table = document.tables[19]
    _replace_cell(qa_table.rows[1].cells[1], "۸۲ آزمون موفق در مجموعهٔ کامل")
    _replace_cell(qa_table.rows[3].cells[1], "outer/inner purge ثبت شد؛ calibration boundary ناقص")
    _replace_cell(qa_table.rows[3].cells[2], "ادعای leakage-free برای v2 رد می‌شود.")

    questions = document.tables[20]
    _replace_cell(
        questions.rows[1].cells[1],
        "خیر؛ هیچ candidate در v2 promote نشد و آزمون تاریخی نیز به ۶۰٪ نرسید.",
    )
    _replace_cell(
        questions.rows[2].cells[1],
        "قابل استنباط نیست؛ v2 leakage مرز calibration دارد و audit تاریخی قبلاً آشکار شده است.",
    )
    _replace_cell(
        questions.rows[3].cells[1],
        "اثبات نشد؛ paired inference معتبر در خانوادهٔ تازه لازم است.",
    )
    _replace_cell(
        questions.rows[4].cells[1],
        "ثبت immutable نتیجه، کشف مستقل نقص و جلوگیری از ادعای unsupported.",
    )
    document.save(output)


def update_report(source: Path, output: Path) -> None:
    shutil.copy2(source, output)
    document = Document(output)
    replacements = {
        84: (
            "در v2 مسیر بازده کامل هر ۱۲ trial ذخیره و PBO/DSR محاسبه شد، اما ممیزی "
            "آماری نشان داد eventهای چندروزه در مرز partitionهای PBO purge نمی‌شوند و DSR "
            "در event time وابستگی‌مقاوم نیست؛ بنابراین این دو سنجه فقط diagnostic هستند."
        ),
        85: "1-7 یافته‌های canonical توسعه v2 و تفسیر آن‌ها",
        86: (
            "هر ۱۲ آزمایش پیش‌ثبت‌شده رد شدند و هیچ candidate به مرحلهٔ freeze، "
            "confirmation، historical audit یا MT5 نرسید."
        ),
        87: (
            "بهترین BA تجمیعی متعلق به no-Hurst/H1 بود: ۰٫۵۲۰۳۵؛ F1 کلان ۰٫۵۱۶۱۸؛ "
            "بازهٔ اطمینان اسمی ۹۵٪ [۰٫۵۰۱۸۹، ۰٫۵۳۸۳۷]. بازده خالص ۵ و ۱۰ bps آن "
            "به‌ترتیب -۰٫۵۹۰۱۱ و -۱٫۱۲۰۷۱ بود."
        ),
        88: (
            "robust-Hurst/H5 بازده نقطه‌ای مثبت داشت، اما معیارهای دقت، پایداری fold، "
            "عدم‌قطعیت، calibration، PBO/DSR و QA را پاس نکرد؛ این نتیجه اثبات سودآوری نیست."
        ),
        89: (
            "هیچ برتری پایدار یا promotion-eligible برای هِرست مشاهده نشد. v2 استنباط "
            "paired مستقیم هِرست در برابر no-Hurst را gate نکرده است؛ ادعای ارزش افزودهٔ "
            "هِرست مجاز نیست."
        ),
        90: (
            "بازهٔ ۲۰۲۲ تا ژوئن ۲۰۲۳ قبلاً دیده شده و confirmation نیست. دورهٔ ژوئیهٔ "
            "۲۰۲۳ تا ژوئیهٔ ۲۰۲۶ historical_audit_previously_revealed است و در v2 بارگذاری نشد."
        ),
        91: (
            "چون candidate وجود ندارد، هیچ MT5 candidate backtest انجام نشد. قرارداد "
            "آینده signal تا order، fill، exit و cost را ردیف‌به‌ردیف تطبیق می‌دهد، اما "
            "native MQL5 inference را ادعا نمی‌کند."
        ),
        92: (
            "ممیزی مستقل پس از freeze نشان داد مرز calibration/evaluation داخلی بر endpoint "
            "برچسب purge نشده است. نبود common-calendar دقیق، familywise correction، "
            "PBO/DSR سازگار با event و provenance بومی نیز باقی است. اصلاح فقط در خانواده‌ای "
            "تازه مجاز است و v2 ادعای leakage-free ندارد."
        ),
    }
    target_paragraphs = {index: document.paragraphs[index] for index in replacements}
    _canonical_notice(document, anchor_index=22, heading_index=49, body_index=52)
    for index, text in replacements.items():
        _replace(target_paragraphs[index], text)
    document.save(output)


def main() -> None:
    root = Path.cwd()
    update_thesis(
        root / "پایان_نامه_HGE_Gold_Forecasting_بازنگری‌شده.docx",
        root / "پایان_نامه_HGE_Gold_Forecasting_canonical_v2.docx",
    )
    update_report(
        root / "گزارش پروژه - بازنگری‌شده.docx",
        root / "گزارش پروژه - canonical_v2.docx",
    )


if __name__ == "__main__":
    main()
